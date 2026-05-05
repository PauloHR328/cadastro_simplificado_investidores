import streamlit as st
import re
from docx import Document
from io import BytesIO
import zipfile
import pandas as pd
import unicodedata


# ==============================
# 🔧 FUNÇÕES
# ==============================

def extrair_variaveis(docx_file):
    from docx import Document
    import re

    doc = Document(docx_file)
    full_text = ""

    def extrair_de_paragrafos(paragraphs):
        text = ""
        for para in paragraphs:
            for run in para.runs:
                text += run.text
            text += "\n"
        return text

    # Parágrafos
    full_text += extrair_de_paragrafos(doc.paragraphs)

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += extrair_de_paragrafos(cell.paragraphs)

    # 🔥 NORMALIZAÇÃO PESADA (ESSENCIAL)
    full_text = full_text.replace('\xa0', ' ')  # espaço invisível
    full_text = re.sub(r"\{\s+\{", "{{", full_text)  # { {
    full_text = re.sub(r"\}\s+\}", "}}", full_text)  # } }
    full_text = re.sub(r"\s+", " ", full_text)  # limpa espaços extras

    # Regex final
    variaveis = re.findall(r"\{\{\s*(.*?)\s*\}\}", full_text)

    return list(set(variaveis))


def sanitizar_nome(nome):
    nome = unicodedata.normalize('NFKD', nome).encode('ASCII', 'ignore').decode('ASCII')
    nome = re.sub(r'[^\w\s-]', '', nome).strip().replace(" ", "_")
    return nome


def renderizar_documento(template_bytes, contexto):
    doc = DocxTemplate(BytesIO(template_bytes))
    doc.render(contexto)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ==============================
# 🚀 APP DE DOCUMENTOS
# ==============================

def render_doc_generator():

    st.markdown("## 📄 Gerador de Documentos Word")
    st.markdown("Upload de template + geração em lote")

    # estado isolado
    if "doc_registros" not in st.session_state:
        st.session_state.doc_registros = []

    uploaded_file = st.file_uploader("Upload do template (.docx)", type=["docx"], key="doc_upload")

    if uploaded_file:
        try:
            template_bytes = uploaded_file.read()

            variaveis = extrair_variaveis(BytesIO(template_bytes))

            if not variaveis:
                st.warning("Nenhuma variável encontrada.")
                return

            st.subheader("Variáveis detectadas")
            st.write(variaveis)

            st.subheader("Novo registro")

            form_data = {}

            for var in variaveis:
                if "data" in var.lower():
                    valor = st.date_input(var, key=f"doc_{var}")
                    if valor:
                        valor = valor.strftime("%d/%m/%Y")
                else:
                    valor = st.text_input(var, key=f"doc_{var}")

                form_data[var] = valor

            if st.button("Adicionar registro", key="add_doc"):

                registro = {
                    k: v if v else "PREENCHIMENTO NECESSÁRIO"
                    for k, v in form_data.items()
                }

                st.session_state.doc_registros.append(registro)
                st.success("Adicionado!")

            # tabela
            if st.session_state.doc_registros:
                st.subheader("Registros")

                df = pd.DataFrame(st.session_state.doc_registros)
                st.dataframe(df)

                idx = st.number_input(
                    "Remover índice",
                    min_value=0,
                    max_value=len(st.session_state.doc_registros) - 1,
                    step=1,
                    key="remove_idx"
                )

                if st.button("Remover"):
                    st.session_state.doc_registros.pop(idx)
                    st.experimental_rerun()

                if st.button("Gerar documentos"):

                    zip_buffer = BytesIO()

                    with zipfile.ZipFile(zip_buffer, "w") as zip_file:

                        for i, registro in enumerate(st.session_state.doc_registros):

                            doc_bytes = renderizar_documento(template_bytes, registro)

                            nome_base = registro.get("nome", f"doc_{i+1}")
                            nome_arquivo = sanitizar_nome(nome_base)

                            zip_file.writestr(f"{nome_arquivo}.docx", doc_bytes.getvalue())

                    zip_buffer.seek(0)

                    st.download_button(
                        "Download ZIP",
                        zip_buffer,
                        "documentos.zip",
                        "application/zip"
                    )

        except Exception as e:
            st.error(f"Erro: {e}")