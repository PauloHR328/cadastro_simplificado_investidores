"""
Utilitários para extração de variáveis e geração de documentos Word em lote.
"""

from __future__ import annotations

import os
import re
import tempfile
import unicodedata
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile


PLACEHOLDER_PATTERN = re.compile(r"{{\s*([^{}]+?)\s*}}")
FILENAME_CANDIDATES = (
    "nome",
    "nome_completo",
    "razao_social",
    "empresa",
    "cnpj",
    "cpf",
)

def is_date_variable(variable_name: str) -> bool:
    """Define automaticamente o tipo do campo com base no nome da variável."""
    normalized_name = variable_name.strip().lower()
    return "data" in normalized_name


def detect_template_kind(template_name: str) -> str:
    """Determina o tipo do template a partir da extensão do arquivo."""
    suffix = Path(template_name).suffix.lower()
    if suffix == ".docx":
        return "word"
    if suffix == ".pdf":
        return "pdf"
    raise ValueError(f"Formato de template não suportado: {suffix or template_name}")


def format_variable_label(variable_name: str) -> str:
    """Converte o nome da variável em um rótulo amigável para a interface."""
    label = re.sub(r"[_-]+", " ", variable_name).strip()
    return label[:1].upper() + label[1:] if label else "Campo"


def _iter_table_paragraphs(table: Any) -> Iterable[str]:
    """Percorre parágrafos dentro das tabelas do documento."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph.text
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _collect_document_texts(document: Any) -> list[str]:
    """Coleta textos do corpo, cabeçalhos, rodapés e tabelas do template."""
    texts: list[str] = []

    texts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        texts.extend(_iter_table_paragraphs(table))

    for section in document.sections:
        header = section.header
        footer = section.footer

        texts.extend(paragraph.text for paragraph in header.paragraphs)
        texts.extend(paragraph.text for paragraph in footer.paragraphs)

        for table in header.tables:
            texts.extend(_iter_table_paragraphs(table))
        for table in footer.tables:
            texts.extend(_iter_table_paragraphs(table))

    return texts


def extract_template_variables(template_bytes: bytes) -> list[str]:
    """Extrai placeholders únicos no formato {{ variavel }} preservando a ordem."""
    from docx import Document

    document = Document(BytesIO(template_bytes))
    found_variables: list[str] = []
    seen_variables: set[str] = set()

    for text in _collect_document_texts(document):
        for match in PLACEHOLDER_PATTERN.findall(text):
            variable = match.strip()
            if variable and variable not in seen_variables:
                seen_variables.add(variable)
                found_variables.append(variable)

    return found_variables


def extract_pdf_form_fields(template_bytes: bytes) -> list[str]:
    """Extrai os nomes dos campos de um PDF editável."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(template_bytes))
    fields = reader.get_fields()
    if not fields:
        return []

    return [field_name.strip() for field_name in fields.keys() if field_name and field_name.strip()]


def extract_pdf_form_field_details(template_bytes: bytes) -> list[dict[str, str]]:
    """Extrai nome e tipo básico dos campos de um PDF editável."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(template_bytes))
    fields = reader.get_fields()
    if not fields:
        return []

    field_details: list[dict[str, str]] = []
    for field_name, field in fields.items():
        cleaned_name = field_name.strip() if field_name else ""
        if not cleaned_name:
            continue

        field_type = str(getattr(field, "field_type", "") or field.get("/FT", ""))
        if field_type == "/Btn":
            normalized_type = "checkbox"
        elif field_type == "/Ch":
            normalized_type = "opcao"
        else:
            normalized_type = "texto"

        field_details.append(
            {
                "name": cleaned_name,
                "field_type": normalized_type,
            }
        )

    return field_details


def extract_template_fields(template_bytes: bytes, template_name: str) -> list[str]:
    """Extrai os campos do template de acordo com o tipo do arquivo."""
    template_kind = detect_template_kind(template_name)
    if template_kind == "word":
        return extract_template_variables(template_bytes)
    if template_kind == "pdf":
        return extract_pdf_form_fields(template_bytes)
    raise ValueError(f"Tipo de template não suportado: {template_kind}")


def extract_template_field_details(
    template_bytes: bytes,
    template_name: str,
) -> list[dict[str, str]]:
    """Extrai detalhes dos campos para apoiar a interface de preenchimento."""
    template_kind = detect_template_kind(template_name)
    if template_kind == "word":
        return [
            {
                "name": variable,
                "field_type": "data" if is_date_variable(variable) else "texto",
            }
            for variable in extract_template_variables(template_bytes)
        ]
    if template_kind == "pdf":
        return extract_pdf_form_field_details(template_bytes)
    raise ValueError(f"Tipo de template não suportado: {template_kind}")


def _format_value(value: Any) -> str:
    """Normaliza o valor para o contexto que será enviado ao template."""
    if value is None:
        return ""

    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")

    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")

    if isinstance(value, str):
        cleaned_value = value.strip()
        return cleaned_value

    return str(value)


def normalize_record(raw_record: dict[str, Any], variables: list[str]) -> dict[str, str]:
    """Garante que todo registro possua todas as variáveis com valor preenchido."""
    return {
        variable: _format_value(raw_record.get(variable))
        for variable in variables
    }


def sanitize_filename(value: str) -> str:
    """Remove caracteres inválidos para nome de arquivo e reduz ruído visual."""
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", ascii_value)
    cleaned = cleaned.strip("._")
    return cleaned[:80] or "documento"


def summarize_record(record: dict[str, str]) -> str:
    """Gera um resumo curto do registro para exibição na interface."""
    for key, value in record.items():
        if value:
            return f"{format_variable_label(key)}: {value}"
    return "Sem identificador principal"


def _pick_record_identifier(record: dict[str, str]) -> str | None:
    """Escolhe o melhor identificador disponível para o nome do arquivo."""
    lowered_key_map = {key.lower(): key for key in record.keys()}

    for candidate in FILENAME_CANDIDATES:
        if candidate in lowered_key_map:
            original_key = lowered_key_map[candidate]
            candidate_value = record.get(original_key)
            if candidate_value:
                return candidate_value

    for value in record.values():
        if value:
            return value

    return None


def build_document_filename(
    template_name: str,
    record: dict[str, str],
    record_index: int,
) -> str:
    """Monta um nome de arquivo amigável e determinístico para cada documento."""
    template_stem = sanitize_filename(Path(template_name).stem)
    template_suffix = Path(template_name).suffix.lower() or ".docx"
    record_identifier = _pick_record_identifier(record)
    suffix = sanitize_filename(record_identifier) if record_identifier else f"registro_{record_index:03d}"
    return f"{template_stem}_{suffix}{template_suffix}"


def build_download_zip_name(template_name: str) -> str:
    """Cria o nome do arquivo zip final usando o template como base."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    template_stem = sanitize_filename(Path(template_name).stem)
    return f"{template_stem}_documentos_{timestamp}.zip"


def render_document(template_bytes: bytes, context: dict[str, str]) -> bytes:
    """Renderiza um documento individual com base no contexto recebido."""
    from docxtpl import DocxTemplate

    temp_input_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_input:
            temp_input.write(template_bytes)
            temp_input_path = temp_input.name

        template = DocxTemplate(temp_input_path)
        template.render(context)

        output_buffer = BytesIO()
        template.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer.getvalue()
    finally:
        if temp_input_path and os.path.exists(temp_input_path):
            os.remove(temp_input_path)


def render_pdf_document(template_bytes: bytes, context: dict[str, str]) -> bytes:
    """Preenche um PDF editável com base no contexto recebido."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(BytesIO(template_bytes))
    writer = PdfWriter()
    writer.append(reader)
    writer.update_page_form_field_values(
        None,
        context,
        auto_regenerate=False,
    )

    output_buffer = BytesIO()
    writer.write(output_buffer)
    output_buffer.seek(0)
    return output_buffer.getvalue()


def generate_documents_zip(
    template_bytes: bytes,
    template_name: str,
    variables: list[str],
    records: list[dict[str, Any]],
) -> tuple[bytes, list[str]]:
    """Gera todos os documentos preenchidos e os compacta em um único zip."""
    template_kind = detect_template_kind(template_name)
    zip_buffer = BytesIO()
    generated_files: list[str] = []
    filename_usage: dict[str, int] = {}

    with ZipFile(zip_buffer, mode="w", compression=ZIP_DEFLATED) as zip_file:
        for index, record in enumerate(records, start=1):
            normalized_record = normalize_record(record, variables)
            if template_kind == "word":
                document_bytes = render_document(template_bytes, normalized_record)
            elif template_kind == "pdf":
                document_bytes = render_pdf_document(template_bytes, normalized_record)
            else:
                raise ValueError(f"Tipo de template não suportado: {template_kind}")

            base_filename = build_document_filename(
                template_name=template_name,
                record=normalized_record,
                record_index=index,
            )

            usage_count = filename_usage.get(base_filename, 0)
            filename_usage[base_filename] = usage_count + 1

            if usage_count:
                stem = Path(base_filename).stem
                suffix = Path(base_filename).suffix
                final_filename = f"{stem}_{usage_count + 1}{suffix}"
            else:
                final_filename = base_filename

            zip_file.writestr(final_filename, document_bytes)
            generated_files.append(final_filename)

    zip_buffer.seek(0)
    return zip_buffer.getvalue(), generated_files
